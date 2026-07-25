"""
Generates a year-over-year income-statement DOCX for financial-statement PDFs.

Vietnamese financial statements (Circular 200/202) lay the income statement out
with two value columns side by side — "Kỳ này" (this period) and "Kỳ trước"
(same period last year) — so the YoY comparison is already in the source PDF;
no historical lookup is needed. The catch: in practice these PDFs are almost
always scanned images with no embedded text layer (confirmed against real
filings from several portfolio companies), so plain text extraction doesn't
work. Instead we OCR the pages with Tesseract (free, local, no external API)
and reconstruct table rows from word bounding boxes, since raw OCR reading
order scrambles multi-column tables.

Best-effort throughout: any failure (Tesseract missing, nothing found) just
skips the report — it never blocks the normal Telegram alert. OCR on dense
Vietnamese number tables isn't perfect, so the generated report carries an
explicit "verify against source" disclaimer.
"""

from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import Optional

import aiohttp
import pdfplumber
from unidecode import unidecode

from config import FIREFOX_HEADERS, FLAG_KEYWORDS, USER_AGENTS
from logger import get_logger

log = get_logger(__name__)

_OUTPUT_DIR = Path(__file__).parent / "reports"
_MAX_PAGES = 15
_MIN_ROWS_TO_STOP = 5  # once this many line items are found, assume the page was the income statement

# Short, robust key phrases — matched by fuzzy word-presence (see _line_matches)
# rather than exact substring, since OCR commonly drops a word or two. Ordered
# to match the standard Circular 200 row sequence top-to-bottom, which lets
# earlier (already-claimed) labels win over later ones on ambiguous rows.
_LINE_ITEMS = [
    ("Doanh thu bán hàng", "Gross revenue"),
    ("giảm trừ doanh thu", "Revenue deductions"),
    ("Doanh thu thuần", "Net revenue"),
    ("Giá vốn hàng bán", "Cost of goods sold"),
    ("Lợi nhuận gộp", "Gross profit"),
    ("Doanh thu hoạt động tài chính", "Financial income"),
    ("Chi phí tài chính", "Financial expenses"),
    ("Chi phí bán hàng", "Selling expenses"),
    ("Chi phí quản lý doanh nghiệp", "G&A expenses"),
    ("Lợi nhuận thuần từ hoạt động kinh doanh", "Operating profit"),
    ("Thu nhập khác", "Other income"),
    ("Chi phí khác", "Other expenses"),
    ("Lợi nhuận khác", "Other profit"),
    ("Tổng lợi nhuận kế toán trước thuế", "Profit before tax"),
    ("Lợi nhuận sau thuế thu nhập doanh nghiệp", "Net profit after tax"),
    ("Lãi cơ bản trên cổ phiếu", "Basic EPS"),
]

_HIGHLIGHT_LABELS = ["Net revenue", "Gross profit", "Operating profit", "Net profit after tax"]

_NUM_RE = re.compile(r"\(?-?[\d.,]{4,}\)?")


def is_financial_statement(item: dict) -> bool:
    """True iff `item` is a PDF flagged with the Financials category."""
    url = (item.get("url") or "").lower()
    if ".pdf" not in url:
        return False
    hay = unidecode(item.get("title") or "").lower()
    return any(
        label == "Financials" and unidecode(kw).lower() in hay
        for kw, label in FLAG_KEYWORDS
    )


async def _download_pdf(url: str) -> bytes:
    from scrapers.base import PlaywrightScraper  # reuse the aiohttp connector builder

    import random

    connector = PlaywrightScraper._build_aiohttp_connector()
    headers = {**FIREFOX_HEADERS, "User-Agent": random.choice(USER_AGENTS)}
    # Scanned financial-statement PDFs run 10-50+ MB; a CI runner's bandwidth
    # to a small VN host can be much slower than local, so give this real room
    # rather than the ~60s that's plenty for ordinary article pages.
    timeout = aiohttp.ClientTimeout(total=180)
    async with aiohttp.ClientSession(headers=headers, connector=connector, timeout=timeout) as session:
        async with session.get(url) as resp:
            resp.raise_for_status()
            return await resp.read()


def _cluster_rows(im, conf_threshold: int = 30, y_tol_frac: float = 0.6) -> list[str]:
    """Reconstruct table rows from OCR word boxes.

    Tesseract's own line grouping frequently splits one visual table row into
    several "lines" when a note-reference column shifts the baseline, which
    scrambles label-to-number association. Clustering by vertical center
    across the whole page (rather than trusting Tesseract's line_num) merges
    those fragments back into complete rows.
    """
    import pytesseract
    from pytesseract import Output

    # Optional override for local dev (e.g. Windows, where Tesseract isn't on
    # PATH). On the GitHub Actions runner, apt-installed tesseract-ocr is on
    # PATH already and neither env var needs to be set.
    if os.getenv("TESSERACT_CMD"):
        pytesseract.pytesseract.tesseract_cmd = os.environ["TESSERACT_CMD"]

    data = pytesseract.image_to_data(im, lang="vie", output_type=Output.DICT)
    words = []
    for i in range(len(data["text"])):
        txt = data["text"][i].strip()
        if not txt:
            continue
        try:
            conf = int(data["conf"][i])
        except (ValueError, TypeError):
            conf = -1
        if conf != -1 and conf < conf_threshold:
            continue
        top, h, left = data["top"][i], data["height"][i], data["left"][i]
        words.append({"top": top, "h": h, "left": left, "text": txt, "center": top + h / 2})

    words.sort(key=lambda w: w["center"])
    rows: list[list[dict]] = []
    for w in words:
        for row in rows:
            avg_h = sum(x["h"] for x in row) / len(row)
            if abs(w["center"] - row[0]["center"]) <= avg_h * y_tol_frac:
                row.append(w)
                break
        else:
            rows.append([w])

    rows.sort(key=lambda row: sum(x["center"] for x in row) / len(row))
    return [" ".join(w["text"] for w in sorted(row, key=lambda w: w["left"])) for row in rows]


def _line_matches(label_vn: str, hay: str) -> bool:
    """Fuzzy match: most of the label's words must appear in the line,
    tolerating a word OCR dropped or garbled."""
    words = unidecode(label_vn).lower().split()
    if not words:
        return False
    hits = sum(1 for w in words if w in hay)
    return hits / len(words) >= 0.7


def _clean_number(raw: str) -> Optional[float]:
    raw = raw.strip()
    neg = raw.startswith("(") and raw.endswith(")")
    raw = raw.strip("()").replace(".", "").replace(",", "")
    if not raw.lstrip("-").isdigit():
        return None
    val = float(raw)
    return -val if neg else val


def _extract_income_statement(pdf_bytes: bytes) -> list[dict]:
    try:
        import pytesseract  # noqa: F401
    except ImportError:
        log.warning("pytesseract not installed — skipping financial report generation")
        return []

    rows_found: dict[str, dict] = {}
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages[:_MAX_PAGES]:
                im = page.to_image(resolution=200).original
                try:
                    row_texts = _cluster_rows(im)
                except Exception as e:
                    log.warning("OCR failed on a page: %s: %s", type(e).__name__, e)
                    continue

                for line in row_texts:
                    hay = unidecode(line).lower()
                    for vn_label, en_label in _LINE_ITEMS:
                        if en_label in rows_found:
                            continue
                        if not _line_matches(vn_label, hay):
                            continue
                        nums = [
                            v for v in (_clean_number(m) for m in _NUM_RE.findall(line))
                            if v is not None and abs(v) >= 1000
                        ]
                        if not nums:
                            continue
                        current = nums[-2] if len(nums) >= 2 else nums[-1]
                        prior = nums[-1] if len(nums) >= 2 else None
                        rows_found[en_label] = {
                            "label": en_label, "label_vn": vn_label,
                            "current": current, "prior": prior,
                        }
                        break

                if len(rows_found) >= _MIN_ROWS_TO_STOP:
                    break
    except Exception as e:
        log.warning("Failed to open/OCR PDF for income statement: %s: %s", type(e).__name__, e)
        return []

    return [
        rows_found[en_label] for _, en_label in _LINE_ITEMS if en_label in rows_found
    ]


# Vietnamese label + narrative-sentence subject to use for each highlighted
# metric — mirrors a sell-side earnings note's bullet style: bolded topic,
# then a plain-language sentence stating the value and its YoY move.
_HIGHLIGHT_VN = {
    "Net revenue": "Doanh thu thuần",
    "Gross profit": "Lợi nhuận gộp",
    "Operating profit": "Lợi nhuận thuần từ hoạt động kinh doanh",
    "Net profit after tax": "Lợi nhuận sau thuế",
}


def _fmt_number(v: Optional[float]) -> str:
    return "—" if v is None else f"{v:,.0f}"


def _fmt_billions_vn(v: Optional[float]) -> str:
    """Format a raw-VND figure in tỷ đồng (billions) with a Vietnamese
    decimal comma, matching how VN equity research quotes figures."""
    if v is None:
        return "—"
    ty = v / 1_000_000_000
    return f"{ty:,.1f}".replace(",", "_").replace(".", ",").replace("_", ".")


def _pct_change(current: Optional[float], prior: Optional[float]) -> Optional[float]:
    if current is None or prior is None or prior == 0:
        return None
    return (current - prior) / abs(prior) * 100


def _fmt_change(current: Optional[float], prior: Optional[float]) -> str:
    pct = _pct_change(current, prior)
    return "—" if pct is None else f"{'+' if pct >= 0 else ''}{pct:.1f}%"


def _trend_word(pct: Optional[float]) -> str:
    if pct is None:
        return "thay đổi"
    return "tăng" if pct >= 0 else "giảm"


def _safe_slug(text: str, maxlen: int = 60) -> str:
    slug = re.sub(r"[^\w\-]+", "_", unidecode(text)).strip("_")
    return slug[:maxlen] or "report"


def _build_headline(company: str, by_label: dict) -> str:
    """A factual, data-driven headline in the same spirit as a sell-side
    note's title — states the two biggest-picture moves, nothing invented."""
    parts = []
    for label in ("Net revenue", "Net profit after tax"):
        r = by_label.get(label)
        if not r or r["current"] is None:
            continue
        pct = _pct_change(r["current"], r["prior"])
        if pct is None:
            continue
        vn = _HIGHLIGHT_VN[label]
        parts.append(f"{vn} {_trend_word(pct)} {abs(pct):.1f}%")
    if not parts:
        return f"{company} — Báo cáo kết quả kinh doanh"
    return f"{company} — {', '.join(parts)} so với cùng kỳ"


def _build_docx(item: dict, rows: list[dict]) -> Path:
    from docx import Document
    from docx.shared import Pt

    _OUTPUT_DIR.mkdir(exist_ok=True)
    by_label = {r["label"]: r for r in rows}

    doc = Document()

    # Header — headline, category/date tag line, source title. Mirrors a
    # sell-side note's header block: entity + headline, "Doanh Nghiệp • date".
    doc.add_heading(_build_headline(item["company"], by_label), level=1)
    tag_line = doc.add_paragraph()
    tag_line.add_run(f"Doanh nghiệp  •  {item.get('published') or '—'}").italic = True
    sub = doc.add_paragraph()
    sub.add_run(item["title"]).italic = True

    # Highlights — narrative sentences for the key lines, each stating the
    # value (in tỷ đồng, as VN equity research conventionally quotes it) and
    # its YoY move, in the same bolded-topic-then-sentence bullet style.
    doc.add_heading("Điểm nhấn", level=2)
    any_highlight = False
    for label in _HIGHLIGHT_LABELS:
        r = by_label.get(label)
        if not r or r["current"] is None:
            continue
        any_highlight = True
        pct = _pct_change(r["current"], r["prior"])
        vn = _HIGHLIGHT_VN[label]
        change_clause = (
            f"({_trend_word(pct)} {abs(pct):.1f}% so với cùng kỳ năm trước)"
            if pct is not None else "(không có số liệu cùng kỳ để so sánh)"
        )
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{vn}: ").bold = True
        p.add_run(f"đạt {_fmt_billions_vn(r['current'])} tỷ đồng {change_clause}.")
    if not any_highlight:
        doc.add_paragraph("Không trích xuất được chỉ tiêu chính nào — xem bảng chi tiết bên dưới.")

    # Full line-item table — this period vs same period last year, the
    # comparison actually printed in a standalone quarterly/annual filing.
    doc.add_heading(f"Hình 1: Kết quả kinh doanh — {item['company']}", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Chỉ tiêu", "Kỳ này (VNĐ)", "Cùng kỳ năm trước (VNĐ)", "Thay đổi YoY"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["label_vn"]
        cells[1].text = _fmt_number(row["current"])
        cells[2].text = _fmt_number(row["prior"])
        cells[3].text = _fmt_change(row["current"], row["prior"])
        if row["label"] in _HIGHLIGHT_LABELS:
            for cell in cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    # Footer — source/note line, matching the "Nguồn: ...; Lưu ý: ..." footnote
    # convention of a sell-side note.
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Nguồn: {item['company']}, {item['url']}. "
        "Lưu ý: Số liệu được trích xuất bằng OCR từ file PDF gốc — vui lòng đối chiếu "
        "với báo cáo tài chính gốc trước khi sử dụng."
    )
    footer_run.font.size = Pt(9)
    footer_run.italic = True

    path = _OUTPUT_DIR / f"{item.get('site_key', 'report')}_{_safe_slug(item['title'])}.docx"
    doc.save(path)
    return path


async def maybe_generate_report(item: dict) -> Optional[Path]:
    """Best-effort: download + OCR + build a YoY income-statement DOCX for a
    financial-statement PDF. Returns the file path, or None if `item` isn't a
    financial-statement PDF, extraction found nothing, or any step failed."""
    if not is_financial_statement(item):
        return None

    try:
        pdf_bytes = await _download_pdf(item["url"])
    except Exception as e:
        # Include the type name — some exceptions (e.g. a bare asyncio.TimeoutError)
        # stringify to "", which would otherwise hide the failure entirely.
        log.warning("[%s] could not download PDF for report: %s: %s", item.get("site_key"), type(e).__name__, e)
        return None

    try:
        rows = _extract_income_statement(pdf_bytes)
    except Exception as e:
        log.warning("[%s] income-statement extraction failed: %s: %s", item.get("site_key"), type(e).__name__, e)
        return None

    if not rows:
        log.info("[%s] no income-statement rows found in %s — skipping DOCX", item.get("site_key"), item["url"])
        return None

    return _build_docx(item, rows)
